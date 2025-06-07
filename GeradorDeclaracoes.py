import os
from docx import Document
from openpyxl import load_workbook
from docx2pdf import convert

def substituir_placeholder(doc, chave, valor):
    for par in doc.paragraphs:
        texto_completo = ''.join(run.text for run in par.runs)
        if chave not in texto_completo:
            continue

        novo_texto = texto_completo.replace(chave, valor)

        # Guardar os estilos originais
        estilos = []
        for run in par.runs:
            estilos.append({
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size
            })

        # Limpar os runs antigos
        for run in par.runs:
            run.clear()
        par._element.clear_content()

        # Adicionar novo run com texto completo e aplicar estilo do primeiro run original
        novo_run = par.add_run(novo_texto)
        if estilos:
            estilo_base = estilos[0]
            novo_run.bold = estilo_base['bold']
            novo_run.italic = estilo_base['italic']
            novo_run.underline = estilo_base['underline']
            novo_run.font.name = estilo_base['font_name']
            novo_run.font.size = estilo_base['font_size']

def substituir_placeholder_mantem_formatacao_mas_nem_sempre_substitui(doc, chave, valor):
    for paragrafo in doc.paragraphs:
        for run in paragrafo.runs:
            if chave in run.text:
                run.text = run.text.replace(chave, valor)

def retirar_espacos_em_branco(s):
    return s.replace(" ", "")

def substituir_nome_no_template(template_path, docx_dir, nome, matricula):
    doc = Document(template_path)
    substituir_placeholder(doc, '{{nome}}', str(nome))
    temp1_path = os.path.join(docx_dir, f"temp_nome_{matricula}.docx")
    doc.save(temp1_path)

    return temp1_path

def substituir_matricula_no_template(temp1_path, docx_dir, nome_sem_espacos, matricula):
    doc2 = Document(temp1_path)
    substituir_placeholder(doc2, '{{matricula}}', str(matricula))
    final_docx_path = os.path.join(docx_dir, f"declaracao_{matricula}_{nome_sem_espacos}.docx")
    doc2.save(final_docx_path)

    return final_docx_path

def converter_docx_para_pdf(pdf_dir, final_docx_path, matricula, nome_sem_espacos):
    pdf_path = os.path.join(pdf_dir, f"declaracao_{matricula}_{nome_sem_espacos}.pdf")
    convert(final_docx_path, pdf_path)
    print(f"\nPDF gerado: {pdf_path}")

def excluir_docx_temporario(temp1_path):
    if os.path.exists(temp1_path):
        os.remove(temp1_path)

def gerar_pdf_para_aluno(template_path, nome, matricula, docx_dir, pdf_dir):
    nome_sem_espacos = retirar_espacos_em_branco(nome)

    # Passo 1: Substituir nome
    temp1_path = substituir_nome_no_template(template_path, docx_dir, nome, matricula)

    # Passo 2: Substituir matrícula
    final_docx_path = substituir_matricula_no_template(temp1_path, docx_dir, nome_sem_espacos, matricula)

    # Passo 3: Converter para PDF
    converter_docx_para_pdf(pdf_dir, final_docx_path, matricula, nome_sem_espacos)

    # Passo 4: Excluir arquivo temporário
    excluir_docx_temporario(temp1_path)

def gerar_declaracoes_em_pdf(template_path, planilha_path, saida_dir):
    wb = load_workbook(planilha_path)
    ws = wb.active

    docx_dir = os.path.join(saida_dir, 'docx')
    pdf_dir = os.path.join(saida_dir, 'pdf')
    os.makedirs(docx_dir, exist_ok=True)
    os.makedirs(pdf_dir, exist_ok=True)

    # Contar o número de alunos (linhas com dados a partir da segunda linha)
    total_alunos = sum(1 for linha in ws.iter_rows(min_row=2, values_only=True) if any(linha)) 
    cont = 0

    for linha in ws.iter_rows(min_row=2, values_only=True):
        matricula, nome = linha

        while True:
            try:
                gerar_pdf_para_aluno(template_path, nome, matricula, docx_dir, pdf_dir)
                break  # Sucesso
            except Exception as e:
                print(f"\nErro ao gerar PDF para {matricula}: {e}")
                print("Tentando novamente...")

        cont += 1

        mostra_barra_progresso(total_alunos, cont, 60, ">")
                
    return total_alunos


def mostra_barra_progresso(total_alunos, cont, tamanho_barra_progresso, caractere_barra_progresso):
    percent_progresso = (100 * cont) / total_alunos    
    barra_progresso_parcial = int((tamanho_barra_progresso * percent_progresso) / 100)

    string_progresso = caractere_barra_progresso * barra_progresso_parcial
    string_falta_completar = " " * (tamanho_barra_progresso - barra_progresso_parcial)
        
    print("[%s%s] = (%d/%d) %.1f%% gerado(s)" % (string_progresso, string_falta_completar, cont, total_alunos, percent_progresso))


# ============================ Exemplo de uso ==========================
template_path = 'modelo_declaracao.docx'
planilha_path = 'alunos.xlsx'
saida_dir = 'declaracoes'

total = gerar_declaracoes_em_pdf(template_path, planilha_path, saida_dir)
print("\nTODAS AS %d DECLARAÇÕES GERADAS COM SUCESSO!" % total)
